#! /usr/bin/env python3

import pdfminer
import docx
import os
from docx.enum.text import WD_COLOR_INDEX
import re
import logging
import traceback
import PyPDF2

from pdfminer.layout import LAParams, LTTextBox, LTTextLine, LTText
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator

caritiveSet = r'with no\b|cariti\w+|abessi\w+|privati\w+|(?<!reproduction prohibited )without|\w+less[\w]*|absen\w+'
keyWords = re.compile(r'{}'.format(caritiveSet), re.IGNORECASE)
toExclude = ['unless', 'regardless', 'nevertheless', 'nonetheless']

listOfPdfs = [file for file in os.listdir(os.getcwd()) if file.endswith('.pdf')]

# Process a pdf file and find pages containing the keywords
def getText(fileName):
    fp = open(fileName, 'rb')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    pages = PDFPage.get_pages(fp)

    listOfPages = list(pages)
        
    founds = list()
    
    for page in listOfPages:
        interpreter.process_page(page)
        layout = device.get_result()
        pageText = ' '.join([lobj.get_text() for lobj in layout if isinstance(lobj, LTTextBox) if lobj.get_text().strip != ''])

        # Erase spaces between each letter in some paragraphs
        tooManySpaces1 = re.compile('(?<=\s\s\w)\s')
        pageText = tooManySpaces1.sub('', pageText)
        tooManySpaces2 = re.compile('(?<=\w\s\w)\s(?!\w\w)|(?<=\w\w)\s(?!\w\w)')
        pageText = tooManySpaces2.sub('', pageText)
        
        # Erase double spaces
        doubleSpace = re.compile('\s{2,}')
        pageText = doubleSpace.sub(' ', pageText)

        # Erase spaces between dots
        spacedDots = re.compile('\.\s\.\s')
        pageText = spacedDots.sub('..', pageText)
        
        # Screen pages with keywords
        if re.search(keyWords, pageText):
            
            #print(re.search(keyWords, pageText).group())
            
            matches, newRegex = exclude(keyWords, pageText)
            if len(matches) != 0 and re.search(newRegex, pageText):
                founds.append("Page {}/{}".format(listOfPages.index(page)+1, len(listOfPages)))
                founds.append(pageText)
    fp.close()

    return founds

def getText2(fileName):
    founds = list()

    fp = open(fileName, 'rb')
    pdf_reader = PyPDF2.PdfFileReader(fp)
    for page in pdf_reader.pages:
        pageText = page.extractText()
        if re.search(keyWords, pageText):
            matches, newRegex = exclude(keyWords, pageText)
            if len(matches) != 0 and re.search(newRegex, pageText):
                founds.append("Page {}/{}".format(list(pdf_reader.pages).index(page)+1, pdf_reader.numPages))
                founds.append(pageText)
    fp.close()
    return founds

# Delete the words to be excluded from the list of matches found on the page
def exclude(keyWords, text):
    matches = re.findall(keyWords, text)
    for match in matches[:]:
        if match.lower() in toExclude:
            matches.remove(match)
    newRegex = re.compile('|'.join([r'{}'.format(match) for match in matches]))
    return matches, newRegex

# Erase repeated pages
#def pageDups(founds):
#    paragraphs = list(dict.fromkeys(founds))
#    return paragraphs

# Add the paragraphs into the doc output file
def buildDoc(founds):
    outputFile = docx.Document()
    for paragraph in founds:
       outputFile.add_paragraph(paragraph)
    return outputFile

def highlight(outputFile):
    # Add highlight in bold to the page numbers and in yellow  to keyword occurrences        
    for i in range(len(outputFile.paragraphs)):
        paragraph = outputFile.paragraphs[i]
        if (i + 1) % 2 == 1:
            paragraph.runs[0].bold = True
        else:
            matches, newRegex = exclude(keyWords, paragraph.text)
            #print("Paragraph {}".format(i+1))
            matches2 = list()
            for match in matches:
                for keyWord in caritiveSet.split('|'):
                    if match.lower() in keyWord:
                        #print(keyWord)
                        matches2.append(keyWord)
                        break
                    else:
                        continue
                if len(matches2) < matches.index(match) + 1:
                    matches2.append(match)
            newRegex = re.compile('|'.join([r'{}'.format(match) for match in matches2]))
            #print(matches)
            #print(matches2)
            #print(newRegex)
            splitParagraph = re.split(newRegex, paragraph.text)
            #print(splitParagraph)
            paragraph.clear()
            for i in range(len(splitParagraph)):
                if i + 1 == len(splitParagraph):
                    paragraph.add_run(splitParagraph[i].strip())
                else:
                    paragraph.add_run(splitParagraph[i].strip() + ' ')
                    paragraph.add_run(matches[i] + ' ').font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    return outputFile

# Count the number of results   
def countResults(outputFile):
    counter = 0
    for paragraph in outputFile.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                counter += 1
    return counter             

# Add the number of results at the beginning of the file
def addCounter(outputFile, counter):
    firstParagraph = outputFile.paragraphs[0]
    firstParagraphText = firstParagraph.text
    firstParagraph.clear()
    if counter == 1:
        firstParagraph.add_run("The search found 1 result.\n\n").bold = True
    else:
        firstParagraph.add_run("The search found {} results.\n\n".format(counter)).bold = True
    firstParagraph.add_run(firstParagraphText).bold = True
    return outputFile

# Launch the processing sequence    
for pdfFile in listOfPdfs:
    print("Processing file {}/{}: {}".format(listOfPdfs.index(pdfFile)+ 1, len(listOfPdfs), pdfFile))
    try:
        founds = getText(pdfFile)
        if len(founds) == 0:
            founds = getText2(pdfFile)
            if len(founds) == 0:
                outputFile = docx.Document()
                outputFile.add_paragraph('Nothing found in this pdf file.')
                outputFile.save('{}_nothingFound.docx'.format(pdfFile[:-4].replace(' ', '_')))
        else:
            outputFile = buildDoc(founds)
            outputFile = highlight(outputFile)
            counter = countResults(outputFile)
            outputFile = addCounter(outputFile, counter)
            outputFile.save('{}_searchResults.docx'.format(pdfFile[:-4].replace(' ', '_')))
    except Exception as e:
        tb = traceback.format_exc()
        with open('{}_error.txt'.format(pdfFile[:-4].replace(' ', '_'), 'w')) as errorLog:
            errorLog.write(tb)
